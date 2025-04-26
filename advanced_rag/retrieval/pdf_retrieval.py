import logging
import os
import time
import io
from typing import Dict, List, Optional, BinaryIO

import fitz
import tempfile
import docx

from advanced_rag.models.document import Document, DocumentMetadata, SourceType
from advanced_rag.utils.text_processor import TextProcessor

logger = logging.getLogger(__name__)

class DocumentRetriever:
    """Component for retrieving content from various document formats."""
    
    def __init__(self, cache_manager=None, text_processor: Optional[TextProcessor] = None):
        self.cache_manager = cache_manager
        self.text_processor = text_processor
    
    def process_document(
        self, 
        file_content: bytes, 
        filename: str,
        content_type: Optional[str] = None,
        query: Optional[str] = None
    ) -> List[Document]:
        """
        Process a document file and extract its content.
        
        Args:
            file_content: Binary content of the file
            filename: Name of the file
            content_type: MIME type of the file
            query: Optional query that triggered this retrieval
            
        Returns:
            List of documents from the file
        """
        try:
            # Determine file type from content_type or filename extension
            file_ext = os.path.splitext(filename)[1].lower()
            
            if content_type == 'application/pdf' or file_ext == '.pdf':
                return self._process_pdf(file_content, filename, query)
            elif (content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' 
                  or file_ext == '.docx'):
                return self._process_docx(file_content, filename, query)
            elif content_type == 'application/msword' or file_ext == '.doc':
                # Note: .doc processing may require additional libraries
                return self._process_doc(file_content, filename, query)
            elif content_type == 'text/plain' or file_ext == '.txt':
                return self._process_text(file_content, filename, query)
            else:
                logger.warning(f"Unsupported file type: {content_type or file_ext}")
                return []
                
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            return []
    
    def _process_pdf(
        self, 
        file_content: bytes, 
        filename: str,
        query: Optional[str] = None
    ) -> List[Document]:
        """Process a PDF file."""
        try:
            # Save the file content to a temporary file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            
            # Open the temporary file with PyMuPDF
            pdf_document = fitz.open(temp_path)
            
            documents = []
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text = page.get_text()
                
                if not text.strip():
                    continue
                    
                content = f"# {filename} (Page {page_num + 1})\n\n{text}"
                
                doc = Document(
                    content=content,
                    metadata=DocumentMetadata(
                        source_type=SourceType.PDF,
                        title=f"{filename} (Page {page_num + 1})",
                        url=None,
                        author=None,
                        publish_date=None,
                        retrieved_date=time.time(),
                        query=query,
                        page_number=page_num + 1,
                    )
                )
                documents.append(doc)
            
            # Clean up the temporary file
            os.unlink(temp_path)
            
            return documents
            
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {e}")
            return []
    
    def _process_docx(
        self, 
        file_content: bytes, 
        filename: str,
        query: Optional[str] = None
    ) -> List[Document]:
        """Process a DOCX file."""
        try:
            # Load the document from bytes
            doc = docx.Document(io.BytesIO(file_content))
            
            # Extract text
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            if not text.strip():
                return []
                
            content = f"# {filename}\n\n{text}"
            
            doc = Document(
                content=content,
                metadata=DocumentMetadata(
                    source_type=SourceType.PDF,  # Using PDF as the source type for all documents
                    title=filename,
                    url=None,
                    author=None,
                    publish_date=None,
                    retrieved_date=time.time(),
                    query=query,
                )
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {e}")
            return []
    
    def _process_doc(
        self, 
        file_content: bytes, 
        filename: str,
        query: Optional[str] = None
    ) -> List[Document]:
        """Process a DOC file."""
        # DOC processing requires additional libraries
        # For simplicity, we'll just extract text if possible
        try:
            # This is a placeholder - for real DOC processing you might need
            # to convert to DOCX first or use a library like textract
            text = "Content could not be extracted from DOC file format."
            
            doc = Document(
                content=f"# {filename}\n\n{text}",
                metadata=DocumentMetadata(
                    source_type=SourceType.PDF,  # Using PDF as the source type
                    title=filename,
                    url=None,
                    author=None,
                    publish_date=None,
                    retrieved_date=time.time(),
                    query=query,
                )
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"Error processing DOC {filename}: {e}")
            return []
    
    def _process_text(
        self, 
        file_content: bytes, 
        filename: str,
        query: Optional[str] = None
    ) -> List[Document]:
        """Process a text file."""
        try:
            # Decode text content
            text = file_content.decode('utf-8', errors='ignore')
            
            if not text.strip():
                return []
                
            content = f"# {filename}\n\n{text}"
            
            doc = Document(
                content=content,
                metadata=DocumentMetadata(
                    source_type=SourceType.PDF,
                    title=filename,
                    url=None,
                    author=None,
                    publish_date=None,
                    retrieved_date=time.time(),
                    query=query,
                )
            )
            
            return [doc]
            
        except Exception as e:
            logger.error(f"Error processing text file {filename}: {e}")
            return []
    
    def process_multiple_documents(
        self, 
        files: List[Dict], 
        query: Optional[str] = None
    ) -> List[Document]:
        """
        Process multiple document files.
        
        Args:
            files: List of dictionaries with file info including content, name, and content_type
            query: Optional query that triggered this retrieval
            
        Returns:
            List of documents from all files
        """
        all_documents = []
        
        for file_info in files:
            file_content = file_info.get("content")
            filename = file_info.get("name")
            content_type = file_info.get("content_type")
            
            if not file_content or not filename:
                continue
                
            documents = self.process_document(file_content, filename, content_type, query)
            all_documents.extend(documents)
        
        return all_documents
